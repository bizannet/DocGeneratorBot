import logging
import json
import os
import re
from datetime import datetime
from pathlib import Path
from weasyprint import HTML, CSS
from jinja2 import Environment, FileSystemLoader, select_autoescape
from config import config
from services.amount_to_words import amount_to_words
from services.file_utils import ensure_dir_exists
from services.document_service import get_template_path as get_service_template_path

logger = logging.getLogger('doc_bot.document_generator')
logger.info("services/document_generator.py ЗАГРУЖЕН УСПЕШНО")


def get_template_path(template_name: str, file_type: str = "autogen") -> Path:
    try:
        logger.info(f"Ищем шаблон: {template_name} (тип файла: {file_type})")
        template_path = get_service_template_path(template_name, file_type)

        if not template_path:
            logger.error(f"Шаблон не найден: {template_name} (тип файла: {file_type})")
            return None

        if not template_path.exists():
            logger.error(f"Файл шаблона не существует: {template_path}")
            return None

        logger.info(f"Найден шаблон: {template_path}")
        return template_path
    except Exception as e:
        logger.error(f"Ошибка при поиске шаблона: {e}", exc_info=True)
        return None


def get_document_path_with_extension(user_id: int, template_name: str, extension: str, doc_name: str = "", suffix: str = "") -> str:
    try:
        user_dir = config.GENERATED_DOCUMENTS_PATH / str(user_id)
        ensure_dir_exists(user_dir)

        if doc_name and suffix:
            safe_doc_name = doc_name.replace('/', '_').replace('\\', '_').replace(':', '_').replace('*', '_').replace('?', '_').replace('"', '_').replace('<', '_').replace('>', '_').replace('|', '_')
            filename = f"{safe_doc_name} {suffix}.{extension}"
        else:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{template_name}_{timestamp}.{extension}"

        full_path = user_dir / filename
        logger.info(f"Сгенерирован путь к документу: {full_path}")
        return str(full_path)
    except Exception as e:
        logger.error(f"Ошибка при генерации пути к документу: {e}", exc_info=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_template_name_backup = template_name.replace('/', '_').replace('\\', '_').replace(':', '_').replace('*', '_').replace('?', '_').replace('"', '_').replace('<', '_').replace('>', '_').replace('|', '_')
        temp_path = config.TEMP_PATH / f"{safe_template_name_backup}_{timestamp}.{extension}"
        logger.info(f"Используем резервный путь к документу: {temp_path}")
        return str(temp_path)


def process_template_data(answers: dict) -> dict:
    logger.info(f"Обработка данных для шаблона. Исходные данные: {answers}")

    template_data = answers.copy()
    template_data['current_date'] = datetime.now()
    template_data['current_date_formatted'] = datetime.now().strftime("%d.%m.%Y")
    template_data['now'] = datetime.now()

    def get_initial(name: str) -> str:
        return (name or "")[:1].upper() if (name and name != "______") else "_"

    template_data["executor_first_name_initial"] = get_initial(template_data.get("executor_first_name"))
    template_data["executor_patronymic_initial"] = get_initial(template_data.get("executor_patronymic"))
    template_data["customer_first_name_initial"] = get_initial(template_data.get("customer_first_name"))
    template_data["customer_patronymic_initial"] = get_initial(template_data.get("customer_patronymic"))

    if 'service_cost_numeric' in template_data:
        try:
            raw = template_data['service_cost_numeric']
            clean = re.sub(r'[^\d.,]', '', str(raw))
            clean = clean.replace(',', '.')
            amount = float(clean)
            template_data['service_cost_full'] = amount_to_words(amount)
        except (ValueError, TypeError, Exception) as e:
            logger.warning(f"Ошибка при обработке суммы '{raw}': {e}")
            template_data['service_cost_full'] = "______ руб. __ коп. (______)"

    for key, value in list(template_data.items()):
        if isinstance(value, list):
            if value and isinstance(value[0], dict):
                template_data[key] = value
            else:
                template_data[key] = ", ".join(str(v) for v in value if v)

    date_fields = [k for k in template_data.keys() if 'date' in k.lower()]
    for field in date_fields:
        if template_data[field] and template_data[field] != "______":
            try:
                for fmt in ["%d.%m.%Y", "%Y-%m-%d", "%Y-%m-%d %H:%M:%S", "%d/%m/%Y"]:
                    try:
                        date_obj = datetime.strptime(template_data[field], fmt)
                        template_data[f"{field}_formatted"] = date_obj.strftime("%d.%m.%Y")
                        break
                    except (ValueError, TypeError):
                        continue
                else:
                    template_data[f"{field}_formatted"] = template_data[field]
            except Exception as e:
                logger.warning(f"Ошибка при форматировании даты {field}: {e}")
                template_data[f"{field}_formatted"] = "______"
        else:
            template_data[f"{field}_formatted"] = "______"

    logger.info(f"Обработанные данные для шаблона: {template_data}")
    return template_data


def date_filter(value, fmt='d.m.Y'):
    try:
        if value == 'now':
            python_fmt = f"%{fmt.replace('.', '.%')}"
            return datetime.now().strftime(python_fmt)
        elif isinstance(value, datetime):
            python_fmt = f"%{fmt.replace('.', '.%')}"
            return value.strftime(python_fmt)
        elif isinstance(value, str):
            for date_fmt in ["%Y-%m-%d", "%d.%m.%Y", "%Y-%m-%d %H:%M:%S"]:
                try:
                    dt = datetime.strptime(value, date_fmt)
                    python_fmt = f"%{fmt.replace('.', '.%')}"
                    return dt.strftime(python_fmt)
                except ValueError:
                    continue
            return str(value)
        else:
            return str(value)
    except Exception as e:
        logger.error(f"Ошибка в фильтре date: {e}")
        return str(value)


def render_template(template_path: Path, answers: dict) -> str:
    try:
        logger.info(f"Начало обработки шаблона: {template_path}")

        template_name = template_path.parent.name
        questions_path = template_path.parent / "questions.json"
        if not questions_path.exists():
            questions_path = Path("documents/templates/contracts") / template_name / "questions.json"
        if not questions_path.exists():
            raise FileNotFoundError(f"Не найден файл вопросов для шаблона: {questions_path}")

        with open(questions_path, 'r', encoding='utf-8') as f:
            questions = json.load(f)

        filled_data = {}
        for idx, question in enumerate(questions):
            key = str(idx)
            step_name = question["step"]
            if key in answers:
                filled_data[step_name] = answers[key]
            else:
                filled_data[step_name] = "______"

        if template_name == "inventory_2025":
            inventory_items = []
            i = 1
            while True:
                name_key = f"item_{i}_name"
                if name_key in answers:
                    name = answers[name_key].strip()
                    if name and name != "______":
                        quantity = answers.get(f"item_{i}_quantity", "______").strip() or "______"
                        condition = answers.get(f"item_{i}_condition", "______").strip() or "______"
                        inventory_items.append({
                            "name": name,
                            "quantity": quantity,
                            "condition": condition
                        })
                        i += 1
                    else:
                        break
                else:
                    break
            filled_data["inventory_items"] = inventory_items

            if "inventory_date" in filled_data and filled_data["inventory_date"] != "______":
                try:
                    d = datetime.strptime(filled_data["inventory_date"], "%d.%m.%Y")
                    months = ["января", "февраля", "марта", "апреля", "мая", "июня",
                              "июля", "августа", "сентября", "октября", "ноября", "декабря"]
                    filled_data["inventory_day"] = str(d.day)
                    filled_data["inventory_month"] = months[d.month - 1]
                    filled_data["inventory_year"] = str(d.year)
                except Exception as e:
                    logger.warning(f"Не удалось распарсить inventory_date: {e}")
                    filled_data["inventory_day"] = "______"
                    filled_data["inventory_month"] = "______"
                    filled_data["inventory_year"] = "______"
            else:
                filled_data["inventory_day"] = "______"
                filled_data["inventory_month"] = "______"
                filled_data["inventory_year"] = "______"

        filled_data = process_template_data(filled_data)
        env = Environment(
            loader=FileSystemLoader(template_path.parent),
            autoescape=select_autoescape(['html', 'xml'])
        )
        env.filters['dateformat'] = lambda dt, fmt='%d.%m.%Y': dt.strftime(fmt) if dt else ''
        env.filters['amount_to_words'] = amount_to_words
        env.filters['date'] = date_filter

        template = env.get_template(template_path.name)
        logger.info(f"Шаблон {template_path.name} успешно загружен")

        html_content = template.render(filled_data)
        logger.info("Шаблон успешно обработан")

        debug_path = config.TEMP_PATH / f"debug_{template_path.stem}.html"
        with open(debug_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        logger.info(f"Отладочный HTML сохранен: {debug_path}")

        return html_content

    except Exception as e:
        logger.error(f"Ошибка при обработке шаблона: {e}", exc_info=True)
        raise

def convert_html_to_pdf(html_content: str, output_path: str) -> bool:
    try:
        logger.info(f"Начало конвертации HTML в PDF. Выходной путь: {output_path}")
        ensure_dir_exists(os.path.dirname(output_path))

        css_path = Path(config.DOCUMENTS_PATH) / "static" / "css" / "document.css"
        logger.info(f"Проверка CSS: {css_path}")

        if css_path.exists():
            logger.info(f"CSS найден: {css_path}")
            css = CSS(filename=str(css_path))
        else:
            logger.warning(f"CSS не найден: {css_path}")
            css = None

        pdf_css = CSS(string='''
            @page {
                size: A4;
                margin: 1.5cm;
            }
            body {
                font-family: Arial, sans-serif;
                line-height: 1.6;
            }
            h1, h2, h3 {
                page-break-after: avoid;
            }
            table {
                page-break-inside: avoid;
            }
            .signature-line {
                border-top: 1px solid black;
                width: 250px;
                text-align: center;
                margin-top: 20px;
            }
        ''')

        HTML(string=html_content).write_pdf(
            target=output_path,
            stylesheets=[css, pdf_css] if css else [pdf_css]
        )

        if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
            logger.info(f"PDF успешно создан: {output_path}")
            return True
        else:
            logger.error(f"PDF файл создан, но пустой или не существует: {output_path}")
            return False

    except Exception as e:
        logger.error(f"Ошибка при конвертации HTML в PDF: {e}", exc_info=True)
        return False


def convert_html_to_docx(html_content: str, output_path: str) -> bool:
    try:
        logger.info(f"Начало конвертации HTML в DOCX. Выходной путь: {output_path}")
        ensure_dir_exists(os.path.dirname(output_path))

        try:
            import mammoth
            result = mammoth.convert_to_docx(html_content)
            result.save(output_path)

            if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
                logger.info(f"DOCX успешно создан через mammoth: {output_path}")
                return True
        except ImportError:
            logger.debug("Библиотека mammoth не установлена. Установите её с помощью 'pip install mammoth' для лучшей конвертации")
        except Exception as e:
            logger.warning(f"Ошибка при конвертации через mammoth: {e}")

        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            from bs4 import BeautifulSoup

            doc = Document()
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(12)

            soup = BeautifulSoup(html_content, 'html.parser')
            body = soup.body or soup

            for element in body.find_all(recursive=False):
                if element.name == 'h1':
                    paragraph = doc.add_heading(level=1)
                    run = paragraph.add_run(element.get_text())
                    run.bold = True
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                elif element.name == 'h2':
                    paragraph = doc.add_heading(level=2)
                    run = paragraph.add_run(element.get_text())
                    run.bold = True
                elif element.name == 'h3':
                    paragraph = doc.add_heading(level=3)
                    run = paragraph.add_run(element.get_text())
                    run.bold = True
                elif element.name == 'p':
                    text = element.get_text().strip()
                    if text:
                        paragraph = doc.add_paragraph(text)
                        if element.get('class') and 'variable' in element.get('class'):
                            for run in paragraph.runs:
                                run.bold = True
                elif element.name == 'ul':
                    for li in element.find_all('li', recursive=False):
                        text = li.get_text().strip()
                        if text:
                            doc.add_paragraph(text, style='List Bullet')
                elif element.name == 'ol':
                    for li in element.find_all('li', recursive=False):
                        text = li.get_text().strip()
                        if text:
                            doc.add_paragraph(text, style='List Number')
                elif element.name == 'div':
                    classes = element.get('class', [])
                    if isinstance(classes, str):
                        classes = [classes]
                    if 'signature' in classes:
                        text = element.get_text().strip()
                        if text:
                            paragraph = doc.add_paragraph()
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                            paragraph.add_run(text)
                    elif 'header' in classes:
                        text = element.get_text().strip()
                        if text:
                            paragraph = doc.add_paragraph()
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            run = paragraph.add_run(text)
                            run.bold = True
                            run.font.size = Pt(14)
                    else:
                        text = element.get_text().strip()
                        if text:
                            doc.add_paragraph(text)
                elif element.string and element.string.strip():
                    doc.add_paragraph(element.string.strip())

            doc.save(output_path)

            if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
                logger.info(f"DOCX успешно создан: {output_path}")
                return True
            else:
                logger.error(f"DOCX файл создан, но пустой или не существует: {output_path}")
                return False

        except ImportError as e:
            if "BeautifulSoup" in str(e):
                logger.error("Библиотека beautifulsoup4 не установлена. Установите её с помощью 'pip install beautifulsoup4'")
            else:
                logger.error("Библиотека python-docx не установлена. Установите её с помощью 'pip install python-docx'")
            return False
        except Exception as e:
            logger.error(f"Ошибка при конвертации через python-docx: {e}", exc_info=True)
            return False

    except Exception as e:
        logger.error(f"Ошибка при конвертации HTML в DOCX: {e}", exc_info=True)
        return False


def generate_document(template_name: str, answers: dict, user_id: int = None, file_type: str = "autogen", doc_name: str = "", suffix: str = "") -> dict:
    logger.info(f"Начата генерация документа: {template_name}. Пользователь: {user_id}")
    logger.info(f"Тип файла для генерации: {file_type}")
    logger.info(f"Данные для генерации: {answers}")
    logger.info(f"Читаемое имя: {doc_name}, суффикс: {suffix}")

    try:
        template_path = get_template_path(template_name, file_type)
        if not template_path:
            logger.error(f"Не удалось найти шаблон: {template_name} (тип файла: {file_type})")
            return None

        html_content = render_template(template_path, answers)

        pdf_path = get_document_path_with_extension(user_id, template_name, 'pdf', doc_name=doc_name, suffix=suffix)
        docx_path = get_document_path_with_extension(user_id, template_name, 'docx', doc_name=doc_name, suffix=suffix)

        logger.info(f"Пути для сохранения: PDF - {pdf_path}, DOCX - {docx_path}")

        pdf_success = convert_html_to_pdf(html_content, pdf_path)
        docx_success = convert_html_to_docx(html_content, docx_path)

        result = {}
        if pdf_success:
            result['pdf'] = str(pdf_path)
        if docx_success:
            result['docx'] = str(docx_path)

        if not result:
            logger.error("Не удалось сгенерировать документы ни в одном формате")
            return None

        logger.info(f"Документы успешно сгенерированы: {result}")
        return result

    except Exception as e:
        logger.error(f"Критическая ошибка при генерации документа: {e}", exc_info=True)
        return None